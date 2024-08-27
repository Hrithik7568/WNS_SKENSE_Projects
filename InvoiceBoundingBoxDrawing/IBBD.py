import streamlit as st
import pdfplumber
import cv2
import numpy as np
from PIL import Image, ImageDraw
from docx import Document
import pytesseract
import re

# Function to extract text and bounding boxes from a PDF
def extract_pdf_info(pdf_path):
    with pdfplumber.open(pdf_path) as pdf:
        first_page = pdf.pages[0]
        text = first_page.extract_text()
        words = first_page.extract_words()
        # Convert the first page to an image
        pil_image = first_page.to_image(resolution=300).original
    return text, words, pil_image

# Function to extract text from a DOCX file and convert it to an image
def extract_docx_info(docx_path):
    doc = Document(docx_path)
    text = "\n".join([para.text for para in doc.paragraphs])

    # Create a blank image
    img = np.ones((1000, 800, 3), np.uint8) * 255

    # Use PIL to draw text onto the image
    pil_img = Image.fromarray(img)
    draw = ImageDraw.Draw(pil_img)
    y = 10
    for para in doc.paragraphs:
        draw.text((10, y), para.text, fill=(0, 0, 0))
        y += 15

    # Convert back to OpenCV image
    img = np.array(pil_img)
    return text, img

# Function to draw bounding boxes around dates, numbers, and percentages using OpenCV and Tesseract
def draw_bounding_boxes(img, text):
    try:
        h, w, _ = img.shape
        boxes = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
        date_pattern = r'\b\d{1,2}\s(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sept|Oct|Nov|Dec)\s\d{4}\b|\b\d{1,2}\s(?:January|February|March|April|May|June|July|August|September|October|November|December)\s\d{4}\b|\b\d{4}[-/]\d{2}[-/]\d{2}\b|\b\d{2}[-/]\d{2}[-/]\d{4}\b|\b\d{2}[-/]\d{2}[-/]\d{4}\b|\b\w{3} \d{2}, \d{4}\b|\b\d{2} \w{3} \d{4}\b'
        number_pattern = r'-?\d+(\.\d+)?'
        percentage_pattern = r'\d+%'

        n_boxes = len(boxes['level'])
        for i in range(n_boxes):
            x, y, bw, bh = boxes['left'][i], boxes['top'][i], boxes['width'][i], boxes['height'][i]
            text_value = boxes['text'][i].lower()

            if re.search(date_pattern, text_value) or re.match(number_pattern, text_value) or re.match(percentage_pattern, text_value):
                color = (0, 0, 255)  # Blue for dates, numbers, and percentages
            else:
                color = (255, 0, 0)  # Red for other matched keywords

            img = cv2.rectangle(img, (x, y), (x + bw, y + bh), color, 2)
        return img
    except Exception as e:
        st.error(f"Error drawing bounding boxes: {e}")
        return img

# Main function to run the Streamlit app
def main():
    st.title("Invoice Bounding Box Drawing")

    uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx"])
    if uploaded_file is not None:
        if uploaded_file.name.endswith('.pdf'):
            text, words, pil_image = extract_pdf_info(uploaded_file)
            st.text_area("Extracted Text", text, height=200)
            image = np.array(pil_image)
            annotated_image = draw_bounding_boxes(image, text)
            if annotated_image is not None:
                # Convert the OpenCV image (annotated_image) to a PIL image
                annotated_image = cv2.cvtColor(annotated_image, cv2.COLOR_BGR2RGB)
                pil_annotated_image = Image.fromarray(annotated_image)
                st.image(pil_annotated_image, caption='Annotated Image', use_column_width=True)

        elif uploaded_file.name.endswith('.docx'):
            text, image = extract_docx_info(uploaded_file)
            st.text_area("Extracted Text", text, height=200)
            annotated_image = draw_bounding_boxes(image, text)
            if annotated_image is not None:
                # Convert the OpenCV image (annotated_image) to a PIL image
                annotated_image = cv2.cvtColor(annotated_image, cv2.COLOR_BGR2RGB)
                pil_annotated_image = Image.fromarray(annotated_image)
                st.image(pil_annotated_image, caption='Annotated Image', use_column_width=True)

if __name__ == "__main__":
    main()

